# backend/app/main.py
import os
import re
from pathlib import Path

from fastapi import FastAPI, APIRouter
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response
from pydantic import BaseModel

# =========================
# .env early load (so quick.py gets env at import time)
# =========================
try:
    from dotenv import load_dotenv
    load_dotenv(dotenv_path=Path(__file__).resolve().parents[1] / ".env", override=False)
except Exception:
    pass

app = FastAPI(title="Pathio Backend")

# =========================
# CORS
# =========================
_env_origins = [o.strip() for o in os.getenv("ALLOWED_ORIGINS", "").split(",") if o.strip()]
default_origins = [
    "https://pathio.streamlit.app",      # Streamlit fallback frontend
    "https://pathio-frontend.onrender.com",  # React production frontend
    "http://localhost:8501",             # local Streamlit dev
    "http://localhost:3000",             # local React/Next.js dev
    "http://localhost:3001",             # local React/Next.js dev (alternative port)
]
allow_origins = _env_origins or default_origins

app.add_middleware(
    CORSMiddleware,
    allow_origins=allow_origins,
    allow_origin_regex=r"^https://[a-zA-Z0-9-]+\.onrender\.com$",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =========================
# CLEAN /export (REGISTERED FIRST)
# =========================
export_router = APIRouter()

class ExportRequest(BaseModel):
    tailored_resume_md: str
    cover_letter_md: str
    which: str  # "resume" or "cover"

def _normalize(s: str) -> str:
    return (s or "").replace("\u00A0", " ").strip()

def _is_header_line(line: str, name: str) -> bool:
    """True if line looks like a header with 'name' (case-insensitive), ignoring *, _, # wrappers."""
    raw = _normalize(line)
    # remove leading header/formatting marks (escape hyphen!)
    raw = re.sub(r"^[#*\s_>\-\.:\|]+", "", raw)
    # remove trailing formatting/space
    raw = re.sub(r"[#*\s_\-\.:\|]+$", "", raw)
    return raw.lower() == name.lower()


def _remove_summary_block(md: str) -> str:
    """Remove a 'Summary' header + its immediate paragraph/bullets."""
    if not md:
        return ""
    lines = md.splitlines()
    out = []
    i = 0
    while i < len(lines):
        if _is_header_line(lines[i], "summary"):
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            while i < len(lines):
                ln = lines[i]
                if not ln.strip():
                    i += 1
                    break
                hdrish = bool(re.match(r"^\s*#+\s+\S", ln)) or bool(re.match(r"^\s*[*_]{1,3}[^*].*[*_]{1,3}\s*$", ln))
                if hdrish:
                    break
                i += 1
            continue
        out.append(lines[i])
        i += 1
    txt = "\n".join(out)
    txt = re.sub(r"\n{3,}", "\n\n", txt).strip()
    return txt

def _remove_what_changed(md: str) -> str:
    """Remove 'What changed' section entirely (from its header to end)."""
    if not md:
        return ""
    lines = md.splitlines()
    for idx, ln in enumerate(lines):
        if _is_header_line(ln, "what changed"):
            return "\n".join(lines[:idx]).rstrip()
    return md

def _strip_markdown_inline(text: str) -> str:
    """Remove markdown marks (handles unbalanced cases like *Heading**)."""
    if not text:
        return ""
    s = text
    s = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", s)  # links
    s = re.sub(r"\*\*(.*?)\*\*", r"\1", s)                 # bold
    s = re.sub(r"__(.*?)__", r"\1", s)                     # bold (alt)
    s = re.sub(r"\*(.*?)\*", r"\1", s)                     # italics
    s = re.sub(r"_(.*?)_", r"\1", s)                       # italics (alt)
    s = re.sub(r"`([^`]+)`", r"\1", s)                     # inline code
    s = re.sub(r"(^|\s)[*_]{1,3}(\S)", r"\1\2", s)         # leading loose *
    s = re.sub(r"(\S)[*_]{1,3}(\s|$)", r"\1\2", s)         # trailing loose *
    s = re.sub(r"[ \t]+", " ", s)                          # collapse spaces
    return s.strip()

def _preprocess_for_resume(md: str) -> str:
    """For résumé export: drop Summary + What changed, then strip inline markdown."""
    if not md:
        return ""
    md2 = _remove_summary_block(md)
    md3 = _remove_what_changed(md2)
    cleaned = "\n".join(_strip_markdown_inline(ln) for ln in md3.splitlines())
    cleaned = re.sub(r"(?m)^\s*-{3,}\s*$", "", cleaned)  # remove hr lines
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned

def _preprocess_for_cover(md: str) -> str:
    """For cover letter export: keep content, just strip inline markdown."""
    if not md:
        return ""
    cleaned = "\n".join(_strip_markdown_inline(ln) for ln in md.splitlines())
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned

@export_router.post("/export")
def export_doc(req: ExportRequest):
    import logging, traceback, io
    logger = logging.getLogger("uvicorn.error")

    try:
        if req.which == "resume":
            content = _preprocess_for_resume(req.tailored_resume_md or "")
        else:
            content = _preprocess_for_cover(req.cover_letter_md or "")

        if not content:
            return Response(
                content=b"No content to export.",
                media_type="text/plain; charset=utf-8",
                headers={
                    "Content-Disposition": 'attachment; filename="export.txt"',
                    "X-Exporter": "clean",
                },
                status_code=400,
            )

        # Try DOCX; fallback to TXT (never label as .docx)
        try:
            from docx import Document
        except Exception:
            return Response(
                content=content.encode("utf-8"),
                media_type="text/plain; charset=utf-8",
                headers={
                    "Content-Disposition": 'attachment; filename="export.txt"',
                    "X-Exporter": "clean",
                    "X-Exporter-Fallback": "no-docx",
                },
                status_code=200,
            )

        doc = Document()
        bullet_re = re.compile(r"^\s*[-*•]\s+(.*)$")

        for raw in content.splitlines():
            line = raw.rstrip("\n")

            if re.match(r"^[A-Z0-9 ,/&()'’.-]{6,}$", line) and line.strip() == line.upper():
                doc.add_heading(line.strip(), level=2)
                continue
            if line.strip().endswith(":") and len(line.strip()) <= 48:
                doc.add_heading(line.strip().rstrip(":"), level=2)
                continue

            m = bullet_re.match(line)
            if m:
                text = m.group(1).strip()
                try:
                    doc.add_paragraph(text, style="List Bullet")
                except Exception:
                    doc.add_paragraph("• " + text)
                continue

            if not line.strip():
                doc.add_paragraph("")
                continue

            doc.add_paragraph(line)

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return Response(
            content=bio.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": 'attachment; filename="pathio_export.docx"',
                "X-Exporter": "clean",
            },
            status_code=200,
        )

    except Exception as e:
        logger.exception("EXPORT_FAILED")
        import traceback as _tb
        debug_text = f"EXPORT_ERROR: {e}\n\nTRACEBACK:\n{_tb.format_exc()}\n"
        return Response(
            content=debug_text.encode("utf-8"),
            media_type="text/plain; charset=utf-8",
            headers={
                "Content-Disposition": 'attachment; filename="export_error.txt"',
                "X-Exporter": "clean",
                "X-Exporter-Error": type(e).__name__,
            },
            status_code=500,
        )

# Register the CLEAN export router FIRST so it takes precedence
app.include_router(export_router)

# =========================
# Other routers (after export)
# =========================
from .routers.quick import router as quick_router
from .routers.jobs import router as jobs_router
try:
    from .routers.public import router as public_router
except Exception:
    public_router = None

app.include_router(quick_router)
app.include_router(jobs_router)
if public_router:
    app.include_router(public_router)

# =========================
# Health & Root
# =========================
@app.get("/healthz")
def healthz():
    return JSONResponse({"ok": True})

@app.get("/")
def root():
    return {
        "service": "Pathio backend",
        "status": "ok",
        "routes": ["/quick-tailor", "/export", "/coach", "/search-jobs", "/analyze-job", "/healthz"],
        "cors_allowed": allow_origins,
    }
