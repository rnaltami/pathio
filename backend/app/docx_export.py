# backend/app/docx_export.py
from docx import Document
from docx.shared import Pt

def export_docx(markdown_text: str, stream):
    """
    Ultra-minimal MD -> DOCX:
    - '# ' becomes bold title
    - '- ' becomes bullet
    - plain lines become paragraphs
    """
    doc = Document()

    for line in (markdown_text or "").splitlines():
        if line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("")  # blank line

    doc.save(stream)
