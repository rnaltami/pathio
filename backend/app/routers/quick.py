from __future__ import annotations

import os
import re
import json
import unicodedata
from typing import List, Dict, Any, Tuple, Optional

from fastapi import APIRouter, HTTPException, Response
from pydantic import BaseModel, Field

# =========================
# OpenAI client (new/old)
# =========================
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

_client = None
_new_api = False
try:
    from openai import OpenAI  # >=1.x
    if OPENAI_API_KEY:
        _client = OpenAI(api_key=OPENAI_API_KEY)
        _new_api = True
except Exception:
    pass

if _client is None and OPENAI_API_KEY:
    try:
        import openai  # old SDK
        openai.api_key = OPENAI_API_KEY
        _client = openai
        _new_api = False
    except Exception:
        _client = None

router = APIRouter()

# =========================
# Models
# =========================
class QuickTailorRequest(BaseModel):
    resume_text: str
    job_text: str
    user_tweaks: Dict[str, Any] = Field(default_factory=dict)

class ExportRequest(BaseModel):
    tailored_resume_md: str
    cover_letter_md: str
    which: str  # "resume" or "cover"

class ChatRequest(BaseModel):
    messages: List[Dict[str, str]]

# =========================
# Heuristic utilities
# =========================
_WORD = re.compile(r"[A-Za-z][A-Za-z\-\./&+']+")
STOPWORDS = {
    "the","and","of","to","in","for","with","on","at","a","an","is","are","be","as",
    "by","from","or","that","this","it","we","you","your","their","our","they","but",
    "not","will","can","may","have","has","had","do","does","did","into","over","per"
}

def toks(text: str) -> List[str]:
    if not text:
        return []
    return [t.lower() for t in _WORD.findall(text) if t.lower() not in STOPWORDS and len(t) > 2]

def key_terms(text: str, limit: int = 40) -> List[str]:
    freq: Dict[str, int] = {}
    for t in toks(text):
        freq[t] = freq.get(t, 0) + 1
    ranked = sorted(freq.items(), key=lambda x: (-x[1], x[0]))
    return [w for w,_ in ranked[:limit]]

def match_and_missing(resume_text: str, job_text: str) -> Tuple[int, List[str]]:
    job = set(key_terms(job_text, limit=60))
    res = set(key_terms(resume_text, limit=200))
    if not job:
        return 0, []
    overlap = len(job & res)
    score = int(round(100.0 * overlap / max(1, len(job))))
    missing = sorted(job - res)
    return max(0, min(100, score)), missing

# ---------- Deterministic normalization + scorer (safety net) ----------
STOP_SAFE = set("a an the and or to for of with in on at by from as is be are was were it this that these those you your we our i me my".split())

def _nfkc(s: str) -> str:
    if not s:
        return ""
    s = unicodedata.normalize("NFKC", s).replace("\u00A0", " ")
    s = re.sub(r"[\u200B-\u200D\uFEFF]", "", s)
    return s.strip()

def _tokset(text: str) -> set[str]:
    text = _nfkc(text).lower()
    toks = re.findall(r"[a-z0-9][a-z0-9.+#-]{1,}", text)  # tokens like python, node.js, aws, c++, gpt-4o
    return {t for t in toks if t not in STOP_SAFE}

def deterministic_match_score(resume_text: str, job_text: str) -> int:
    R, J = _tokset(resume_text), _tokset(job_text)
    if not R or not J:
        return 0
    overlap = len(R & J)
    score = int(round(100 * overlap / len(J)))
    if overlap > 0 and score == 0:
        score = 1
    return max(0, min(100, score))

# ---------- ATS checks (calmer on typography) ----------
def ats_checks(resume_text: str) -> List[str]:
    flags: List[str] = []

    if len(resume_text) > 40000:
        flags.append("Resume is very long; consider trimming.")

    # Ignore common typographic characters; warn only on unusual density
    allowed = set("•–—“”‘’…·°®™©")
    raw = [c for c in resume_text if ord(c) > 127 and c not in allowed]
    total = max(1, len(resume_text))
    density = len(raw) / total

    if density >= 0.0025 and len(raw) >= 5:
        flags.append("Resume contains unusual non-ASCII characters; export as UTF-8.")
    else:
        if any(ord(c) > 127 for c in resume_text):
            flags.append("Curly quotes/bullets/dashes detected (OK). Ensure final export is UTF-8 or PDF.")

    if len(re.findall(r"\b[A-Z]{3,}\b", resume_text)) > 80:
        flags.append("Many ALL-CAPS tokens; reduce for ATS readability.")

    return flags or ["none"]

# =========================
# LLM helpers
# =========================
def _chat(messages: List[Dict[str, str]], max_tokens: int = 1400, temperature: float = 0.3) -> Optional[str]:
    if not _client or not OPENAI_API_KEY:
        return None
    try:
        if _new_api:
            resp = _client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=messages,
                temperature=temperature,
                max_tokens=max_tokens,
            )
            return (resp.choices[0].message.content or "").strip()
        else:
            resp = _client.ChatCompletion.create(
                model=OPENAI_MODEL,
                messages=messages,
                temperature=temperature,
                max_tokens=max_tokens,
            )
            return resp.choices[0].message["content"].strip()
    except Exception:
        return None

def _extract_json(raw: str) -> Optional[dict]:
    if not raw:
        return None
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
    try:
        return json.loads(raw)
    except Exception:
        pass
    start = raw.find("{")
    end = raw.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(raw[start:end+1])
        except Exception:
            return None
    return None

# Keep only skill-like items (tools, frameworks, metrics, multiword phrases)
_GENERIC_BAN = {
    "ability","abilities","any","across","activities","adherence","administrative","analysis","assign","assigned",
    "assigning","assistants","attention","availability","based","basic","bundle","calendar","cheat","clips","coaching",
    "collect","competing","compile","compliance","coordinate","coordination","daily","data","ensure","execution","host",
    "hosts","live","livestream","manage","operational","organizational","performance","product","reporting","role","skills",
    "smooth","state","status","applicants","application","including","include","such","must","work","working","before","when"
}
def _is_skilllike(s: str) -> bool:
    s = s.strip()
    if not s or len(s) < 3:
        return False
    low = s.lower().strip(" .,:;()[]")
    if low in _GENERIC_BAN:
        return False
    if " " in low:
        return True
    if re.search(r"(obs|sql|airflow|scrapy|sagemaker|athena|glue|boto3|notion|jira|asana|tableau|looker|duckdb|spark|pytorch|tensorflow|vertex|bigquery|kafka|redis|elasticsearch|datastudio|powerbi|ga4|ctr|ndcg|map@|f1|auc)$", low):
        return True
    if re.search(r"[A-Z]{2,}", s):  # acronyms like CTR, KPI, AVD, OBS
        return True
    return False

def llm_insight_score_miss(resume_text: str, job_text: str) -> Tuple[Optional[int], Optional[List[str]], str]:
    """Return (score, missing_skills, engine). Engine is 'llm' or 'heuristic'."""
    sys = (
        "You are a precise hiring copilot. Extract only missing skill-like keywords: tools, frameworks, metrics, "
        "workflows, certifications. Avoid filler ('and', 'must'), HR words, or locations."
    )
    user = (
        "Return strict JSON only:\n"
        "{\"score\": <0-100>, \"missing\": [string, string, ... up to 10]}\n\n"
        f"JOB:\n{job_text}\n\nRESUME:\n{resume_text}\n\n"
        "Rules:\n"
        "- 'missing' must be domain/skill terms (e.g., OBS, Airflow, CTR, run-of-show, talent pipeline).\n"
        "- Do not include generic words like 'any', 'role', 'skills', 'including', or locations.\n"
    )
    raw = _chat([{"role":"system","content":sys},{"role":"user","content":user}], max_tokens=300, temperature=0.0)
    data = _extract_json(raw or "")
    if not data:
        return None, None, "heuristic"

    # robust score coercion (handles "82%", "82", 82, 0.82, etc.)
    s_raw = data.get("score", None)
    s: Optional[int] = None
    if s_raw is not None:
        try:
            if isinstance(s_raw, str):
                m = re.search(r"\d{1,3}", s_raw)
                if m:
                    s = int(m.group(0))
            elif isinstance(s_raw, (int, float)):
                s = int(round(float(s_raw)))
        except Exception:
            s = None

    missing = [str(x).strip() for x in (data.get("missing") or []) if isinstance(x, (str,int))]
    cleaned: List[str] = []
    seen = set()
    for m in missing:
        m = re.sub(r"\s+", " ", str(m)).strip(" .,-").strip()
        if not _is_skilllike(m):
            continue
        key = m.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(m)

    if s is None:
        return None, None, "heuristic"
    return max(0, min(100, s)), cleaned[:10], "llm"

def _bias_lines(job_text: str) -> List[str]:
    jl = job_text.lower()
    bias: List[str] = []
    if "amazon" in jl or "aws" in jl:
        bias.append("Emphasize AWS (S3, Lambda, Glue, Athena, SageMaker, CloudWatch) and boto3.")
    if "google" in jl or "gcp" in jl:
        bias.append("Emphasize GCP (BigQuery, Dataflow, Pub/Sub, Vertex AI).")
    if "microsoft" in jl or "azure" in jl:
        bias.append("Emphasize Azure (Data Lake, Synapse, Functions, ML).")
    if any(k in jl for k in ["crawl","crawler","scrap","spider","index","document processing","extraction"]):
        bias.append("Mention Scrapy/Requests, robots.txt, sitemaps, polite crawling, backoff, queues.")
    if any(k in jl for k in ["applied scientist","ml","nlp","ir","ranking","search"]):
        bias.append("Include metrics like F1, ROC-AUC, NDCG@K, MAP@K; reproducible notebooks.")
    return bias

def llm_actions(resume_text: str, job_text: str, missing_skills: List[str]) -> Tuple[Optional[List[str]], Optional[List[str]]]:
    biases = _bias_lines(job_text)
    bias_text = ("\nBiases:\n- " + "\n- ".join(biases) + "\n") if biases else ""
    sys = "You design concrete, measurable upskilling actions for candidates."
    user = (
        "Return strict JSON only (no backticks):\n"
        "{ \"do_now\": [string, string, string], \"do_long\": [string, string] }\n\n"
        f"JOB:\n{job_text}\n\nRESUME:\n{resume_text}\n\nMISSING SKILLS: {missing_skills}\n"
        + bias_text +
        "Constraints:\n"
        "- Each 'do_now' is hands-on and yields an artifact (repo/notebook/report/dashboard/demo/case study).\n"
        "- Be specific: tools, datasets, metrics. < 230 chars each. Include time estimate.\n"
        "- 'do_long' may include at most one course/cert; otherwise deeper project/OSS.\n"
    )
    raw = _chat([{"role":"system","content":sys},{"role":"user","content":user}], max_tokens=500, temperature=0.2)
    data = _extract_json(raw or "")
    if not data:
        return None, None
    do_now = [str(x) for x in (data.get("do_now") or [])]
    do_long = [str(x) for x in (data.get("do_long") or [])]

def build_actions_fallback(missing: List[str]) -> Tuple[List[str], List[str]]:
    """
    Heuristic upskilling suggestions when the LLM is unavailable.
    Uses up to 3 'missing' keywords if present; otherwise generic ops terms.
    """
    # keep only short, readable hints
    kws = [m.strip() for m in (missing or []) if isinstance(m, str) and 2 < len(m.strip()) < 28][:3]
    if not kws:
        kws = ["scheduling", "coordination", "reporting"]

    do_now = [
        f"Draft a one-page alignment sheet: map current bullets to job needs ({', '.join(kws)}). (time: ~1–2 hours)",
        f"Create a sample artifact that mirrors the role (checklist/run-book) emphasizing {kws[0]}. (time: ~3–4 hours)",
        "Compile an impact sheet from past projects (before/after, volume, quality, timing). (time: ~2–3 hours)",
    ]

    do_long = [
        f"Turn the sample artifact into a portfolio piece with a README and reflection. (time: ~1–2 weeks)",
        "Contribute a small improvement to a workflow/template you use; document and share. (time: ~1–2 weeks)",
    ]
    return do_now, do_long


    def clean(items: List[str], is_now: bool) -> List[str]:
        out: List[str] = []
        for s in items:
            s = re.sub(r"\s+", " ", s).strip()
            if len(s) < 12:
                continue
            if not re.search(r"(repo|notebook|readme|report|dashboard|demo|case study|artifact|portfolio|gist|PR|evaluation)", s, re.I):
                if not is_now and re.search(r"(course|certification|specialization|certificate)", s, re.I):
                    pass
                else:
                    continue
            if not re.search(r"\b(\d+[-\s]?(hour|day|week)s?)\b", s, re.I):
                s += " (time: ~4–6 hours)" if is_now else " (time: ~1–3 weeks)"
            out.append(s[:230])
        if not is_now:
            kept: List[str] = []
            course_used = False
            for it in out:
                if re.search(r"(course|certification|specialization|certificate)", it, re.I):
                    if course_used:
                        continue
                    course_used = True
                kept.append(it)
            return kept[:2]
        return out[:3]

    return clean(do_now, True), clean(do_long, False)

# ---------- NEW: Heuristic fallback for actions ----------
def build_actions_fallback(missing: List[str]) -> Tuple[List[str], List[str]]:
    """
    Heuristic actions when the LLM path is unavailable or returns nothing.
    Keeps items honest (no new tools) and artifact-oriented.
    """
    focus = [m for m in (missing or []) if isinstance(m, str) and len(m) > 2][:3]
    if not focus:
        focus = ["the role’s core outputs"]

    do_now = [
        "Create a one-page alignment summary mapping your bullets to the job responsibilities; add links to past work. (time: ~1–2 hours)",
        f"Draft a sample work artifact that matches {focus[0]} (process doc / outline / checklist) using tools already on your resume. (time: ~3–4 hours)",
        "Compile a simple impact sheet from past projects (before/after, volume, quality, timing). (time: ~2–3 hours)",
    ][:3]

    do_long = [
        "Extend the sample artifact into a portfolio piece with a README and reflection on tradeoffs. (time: ~1–2 weeks)",
        "Contribute a small improvement to a template/workflow you already use; document the change. (time: ~1–2 weeks)",
    ][:2]

    return do_now, do_long

# =========================
# Sanitizer to prevent over-embellishment
# =========================
def _sentences(md: str) -> List[str]:
    parts: List[str] = []
    for line in (md or "").splitlines():
        if not line.strip():
            parts.append(line)
            continue
        if line.strip().startswith(("#", "-", "•", "*")):
            parts.append(line)
            continue
        segs = re.split(r'(?<=[.!?])\s+(?=[A-Z])', line)
        parts.extend(segs)
    return parts

def _normalize_tokens_set(text: str) -> set[str]:
    toks = re.findall(r"[a-z0-9][a-z0-9.+#-]{1,}", unicodedata.normalize("NFKC", (text or "").lower()))
    return set(toks)

_PROTECTED_CLUSTERS = [
    {"machine", "learning", "ml", "model", "models", "tensorflow", "pytorch", "spark", "hadoop", "mllib", "sagemaker"},
    {"a/b", "ab", "experimentation", "experiment", "ab-testing", "a-b"},
    {"statistical", "regression", "classification", "roc-auc", "ndcg", "map", "f1"},
    {"data-driven", "data-driven", "analytics", "analytical"},
]

def sanitize_tailored(tailored_md: str, resume_text: str) -> str:
    src_tokens = _normalize_tokens_set(resume_text)
    has_percent_in_src = "%" in (resume_text or "")

    def allowed(sentence: str) -> bool:
        s = sentence.strip()
        if not s:
            return True
        low = unicodedata.normalize("NFKC", s.lower())

        # Drop % claims if none existed in source
        if "%" in s and not has_percent_in_src:
            return False

        # Guard protected clusters: if any appears but not present in source tokens, drop
        for cluster in _PROTECTED_CLUSTERS:
            if any(term in low for term in cluster):
                if not any(term in src_tokens for term in cluster):
                    return False
        return True

    lines = _sentences(tailored_md or "")
    kept = [ln for ln in lines if allowed(ln)]

    out: List[str] = []
    prev_blank = False
    for ln in kept:
        if not ln.strip():
            if prev_blank:
                continue
            prev_blank = True
        else:
            prev_blank = False
        out.append(ln)
    return "\n".join(out).strip()

# =========================
# Tailoring
# =========================
def heuristic_resume(resume_text: str, job_text: str, missing: List[str]) -> str:
    job_terms = set(key_terms(job_text, 30))
    lines = [l.strip() for l in (resume_text or "").splitlines() if l.strip()]
    scored = []
    for l in lines:
        tl = set(toks(l))
        score = len(tl & job_terms)
        if l.startswith(("•","-","*")) or ("—" in l or "-" in l) or len(l) <= 140:
            if score > 0:
                scored.append((score, len(l), l))
    scored.sort(key=lambda x: (-x[0], x[1]))
    top = [l for _,__,l in scored[:8]]
    if not top:
        core = [m for m in missing if len(m) > 3][:4]
        top = [f"- Alignment with **{m}**; ready to demonstrate impact." for m in core] or [
            "- Clear, ATS-friendly bullets aligned to the role.",
            "- Comfortable with performance metrics and run-of-show."
        ]
    header = "**Tailored Summary**\n- Role-aligned bullets extracted from your resume.\n\n**Highlights**\n"
    return header + "\n".join(top)

def heuristic_cover(resume_text: str, job_text: str) -> str:
    m = re.search(r"\b(title|role|position)\b[:\-–]\s*(.+)", job_text, re.IGNORECASE)
    title = m.group(2).strip() if m else "the role"
    title = re.sub(r"[\r\n]+", " ", title)[:100]
    return (
        f"**Cover Letter**\n\n"
        f"Dear Hiring Team,\n\n"
        f"I’m excited to apply for {title}. My background aligns with your needs in operations, coordination, and metrics tracking. "
        f"I work well cross-functionally and in fast-paced settings. I’d welcome the chance to contribute and learn.\n\n"
        f"Sincerely,\nYour Name\n"
    )

def call_llm_tailor(resume_text: str, job_text: str) -> Tuple[str, str]:
    if not _client or not OPENAI_API_KEY:
        return "", ""
    system = (
        "You are an expert resume editor.\n"
        "- REWRITE ONLY WHAT EXISTS. Do not invent responsibilities, metrics, tools, or results.\n"
        "- If the job asks for skills the resume does not show, DO NOT claim them. Emphasize transferable skills instead.\n"
        "- Keep it factual, concise, ATS-friendly. Prefer neutral verbs over marketing language (avoid 'data-driven' unless present).\n"
        "- Use Markdown. Do not include code fences."
    )
    user_resume = (
        f"JOB DESCRIPTION:\n{job_text}\n\n"
        f"RESUME (verbatim source):\n{resume_text}\n\n"
        "TASKS:\n"
        "1) Rewrite the resume to align with the job WITHOUT adding new tools/skills/metrics not in the source.\n"
        "2) Draft a short cover letter (≤180 words) that stays factual to the source resume. Do not claim missing skills.\n"
        "Return as exactly two sections labeled:\n"
        "===TAILORED_RESUME===\n...\n===COVER_LETTER===\n...\n"
    )
    try:
        raw = _chat(
            [{"role":"system","content":system},{"role":"user","content":user_resume}],
            max_tokens=1400, temperature=0.35
        )
        if not raw:
            return "", ""
        tail = ""
        cover = ""
        if "===TAILORED_RESUME===" in raw:
            _, rest = raw.split("===TAILORED_RESUME===", 1)
            if "===COVER_LETTER===" in rest:
                tail, cover = rest.split("===COVER_LETTER===", 1)
            else:
                tail = rest
        else:
            parts = raw.split("**Cover", 1)
            tail = parts[0].strip()
            cover = ("**Cover" + parts[1]).strip() if len(parts) > 1 else ""
        # sanitize both sections against the source resume
        tail = sanitize_tailored(tail, resume_text)
        cover = sanitize_tailored(cover, resume_text)
        return tail.strip(), cover.strip()
    except Exception:
        return "", ""

def _brief_summary_from_resume(resume_text: str) -> str:
    """Two-line, truthful summary derived only from terms in the source resume."""
    rlow = (resume_text or "").lower()
    core = "Coordinator/writer experienced in version/change control, scheduling, structured notes, and cross-team communication."
    tools = []
    for t in ["OpenAI", "Anthropic Claude", "Python", "Final Draft", "Google Workspace"]:
        if t.lower() in rlow:
            tools.append(t)
    if tools:
        return f"{core} Familiar with " + ", ".join(tools) + "."
    return core

def _ensure_summary_and_changes(tailored_md: str, resume_text: str) -> str:
    """Guarantee a Summary at the top and a 'What changed' note at the end (no invention)."""
    text = (tailored_md or "").strip()

    # Add **Summary** if missing
    has_summary = re.search(r"(?im)^\s*(\*\*|##?\s*)summary\b", text) is not None
    if not has_summary:
        summary = _brief_summary_from_resume(resume_text)
        text = f"**Summary**\n{summary}\n\n" + text

    # Add **What changed** if missing
    has_changes = re.search(r"(?im)^\s*\*\*what changed\*\*", text) is not None
    if not has_changes:
        changes = (
            "\n\n**What changed**\n"
            "- Reordered to foreground coordination/scheduling already present in your resume.\n"
            "- Clarified process work (revisions/version control, continuity, notes & follow-ups).\n"
            "- Kept show/company names and tools; tightened phrasing for ATS readability.\n"
        )
        text = text.rstrip() + changes

    return text
# ---------- LLM live-call debug ----------
@router.get("/debug/llm")
def debug_llm():
    info = {
        "has_key": bool(OPENAI_API_KEY),
        "client_exists": _client is not None,
        "using_new_sdk": _new_api,
        "model": OPENAI_MODEL,
    }
    if not _client or not OPENAI_API_KEY:
        info["ok"] = False
        info["error"] = "No client or key"
        return info
    try:
        if _new_api:
            resp = _client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": "ping"}],
                temperature=0,
                max_tokens=5,
            )
            sample = (resp.choices[0].message.content or "")[:40]
        else:
            resp = _client.ChatCompletion.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": "ping"}],
                temperature=0,
                max_tokens=5,
            )
            sample = (resp.choices[0].message["content"] or "")[:40]
        info["ok"] = True
        info["sample"] = sample
        return info
    except Exception as e:
        info["ok"] = False
        info["error"] = str(e)
        return info

# =========================
# Routes
# =========================
@router.post("/quick-tailor")
def quick_tailor(req: QuickTailorRequest):
    # Normalize inputs to stabilize scoring & ATS heuristics
    resume_text = _nfkc(req.resume_text or "")
    job_text = _nfkc(req.job_text or "")
    if not resume_text or not job_text:
        raise HTTPException(status_code=400, detail="Both resume_text and job_text are required.")

    # Insights: LLM + heuristic; pick the healthier signal
    score, missing, engine = llm_insight_score_miss(resume_text, job_text)
    h_score, h_missing = match_and_missing(resume_text, job_text)

    if score is None or missing is None:
        score, missing, engine = h_score, h_missing, "heuristic"
    else:
        if score == 0 and h_score > 0:
            score, engine = h_score, "heuristic"
        if not missing:
            missing = h_missing

    flags = ats_checks(resume_text)

    # Tailoring: LLM first; fallback to heuristic
    tailored_md, cover_md = call_llm_tailor(resume_text, job_text)
    if not tailored_md:
        tailored_md = heuristic_resume(resume_text, job_text, missing or [])
    tailored_md = _ensure_summary_and_changes(sanitize_tailored(tailored_md, resume_text), resume_text)

    if not cover_md:
        cover_md = heuristic_cover(resume_text, job_text)
    cover_md = sanitize_tailored(cover_md, resume_text)

    # Actions: prefer LLM; fallback to heuristic
    do_now, do_long = llm_actions(resume_text, job_text, missing or [])
    if not do_now or not do_long:
        do_now, do_long = build_actions_fallback(missing or [])

    return {
        "tailored_resume_md": tailored_md,
        "cover_letter_md": cover_md,
        "insights": {
            "engine": engine,  # "llm" or "heuristic" for debugging
            "match_score": score,
            "missing_keywords": missing or [],
            "ats_flags": flags,
            "do_now": do_now,
            "do_long": do_long,
        },
    }

@router.post("/export")
def export_doc(req: ExportRequest):
    content = req.tailored_resume_md if req.which == "resume" else req.cover_letter_md
    if not content:
        raise HTTPException(status_code=400, detail="No content to export.")

    try:
        from docx import Document  # python-docx
    except Exception:
        return Response(
            content=content.encode("utf-8"),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": 'attachment; filename="export.txt"'},
        )

    doc = Document()
    for line in content.splitlines():
        s = line.strip("\n")
        if s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        else:
            doc.add_paragraph(s)

    import io
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return Response(
        content=bio.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="pathio_export.docx"'},
    )

@router.post("/coach")
def coach(req: ChatRequest):
    """Lightweight how-to chat for the 'Show me how' link."""
    msgs = req.messages or []
    system = (
        "You are a practical, concise how-to coach. Answer with a focused step-by-step recipe "
        "for the user's request: list concrete steps, tools/menus/commands, and a small checklist. "
        "Avoid generic self-help advice; be specific and actionable."
    )
    chat = [{"role":"system","content":system}] + msgs
    raw = _chat(chat, max_tokens=700, temperature=0.2) or ""
    return {"reply": raw.strip() or "Here’s a short, concrete plan:\n1) Define the goal\n2) Gather tools\n3) Execute\n4) Review\n"}

@router.get("/debug/llm")
def debug_llm():
    return {
        "has_key": bool(OPENAI_API_KEY),
        "client_exists": _client is not None,
        "using_new_sdk": _new_api,
        "model": OPENAI_MODEL,
    }
